import os
from pandas import read_excel
import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import datetime
import glob
import tkinter
from tkinter import filedialog
from sys import exit

class Members():
    listFinishNumber = None
    listName = None
    listBirth = None
    document = None
    paras = None
    countCertification = 1

    def __init__(self, path):
        self.document = Document('./data/certForm.docx')
        eduMembers = read_excel(path)
        self.listFinishNumber = eduMembers['이수번호']
        self.listName = eduMembers['성명']
        self.listBirth = eduMembers['생년월일']
        self.paras = self.document.paragraphs

    def enterData(self, count):
        #이수번호
        finish_number = '이수번호     ' +self.listFinishNumber[count] + '-' + str(Members.countCertification)
        Members.countCertification = Members.countCertification + 1
        self.paras[8].text = ''
        run = self.paras[8].add_run(finish_number)
        run.font.name = '휴먼명조'
        run.font.size = Pt(14)
        run.font.bold = True

        #성명, 생년월일
        name_and_birth = '성    명     ' + self.listName[count] +'                    생년월일 '+ self.listBirth[count]
        self.paras[9].text = ''
        run = self.paras[9].add_run(name_and_birth)
        run.font.name = '휴먼명조'
        run.font.size = Pt(14)
        run.font.bold = True

        #오늘 날짜
        currnetDay = datetime.date.today()
        curYear = currnetDay.year
        curMonth = currnetDay.month
        curDay = currnetDay.day
        #발급날짜(오늘)
        textDate = '{0} 년   {1} 월   {2} 일'.format(curYear, curMonth, curDay)
        self.paras[19].text = ''
        run = self.paras[19].add_run(textDate)
        run.font.name = '휴먼명조'
        run.font.size = Pt(14)
        run.font.bold = True
        self.paras[19].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def saveDocument(self, count):
        save_file = './data/'+str(datetime.date.today())+'/'+str(datetime.date.today().year)+str(datetime.date.today().month)+str(datetime.date.today().day)+str(self.listName[count])+' 수료증.docx'
        self.document.save(save_file)

    def getCountMembers(self):
        return len(self.listName)


#문서 통합하기
def merge_certification():
    filesPath = './data/'+str(datetime.date.today())+'/'+str(datetime.date.today().year)+str(datetime.date.today().month)+str(datetime.date.today().day)+'*수료증.docx'
    files = glob.glob(filesPath)
    print(files)
    merged_document = Document('./data/default.docx')
    for index, file in enumerate(files):
        target_doc = Document(file)

        if index < len(files)-1:
            target_doc.add_page_break()

        for element in target_doc.element.body:
            merged_document.element.body.append(element)

        merged_document.save('./data/'+str(datetime.date.today())+'/통합수료증/'+str(datetime.date.today().year)+str(datetime.date.today().month)+str(datetime.date.today().day)+'통합수료증.docx')


#수료증 저장될 위치 생성하기
def make_directory(dir):
    try:
        if not(os.path.isdir(dir)):
            os.makedirs(os.path.join(dir))
    except OSError as e:
        if e.errno != e.errno.EEXIST:
            print("failed to make directory")
            raise


def cert_generator():
    excel_cert_path = filedialog.askopenfilename(initialdir='./data/', filetypes=[("Excel files", "*.xlsx")], title='파일 탐색기')
    if excel_cert_path:
        gernerate_direction = './data/'+str(datetime.date.today())+'/통합수료증'
        cert_members = Members(excel_cert_path)
        memberCount = int(cert_members.getCountMembers())
        make_directory(gernerate_direction)
        for count in range(0, memberCount):
            cert_members.enterData(count)
            cert_members.saveDocument(count)
        merge_certification()
    else :
        print("다시선택하쇼")


if __name__ ==  "__main__":
    window = tkinter.Tk()

    window.title("수료증 생성기")
    window.geometry("600x100")
    window.resizable(False, False)

    text = "수료증 생성기입니다. 생성하기 버튼을 눌러서 수료대상 엑셀파일을 지정하면 자동으로 생성됩니다"
    label1 = tkinter.Label(window, text=text)
    label1.pack()
    label2 = tkinter.Label(window)
    label2.pack()

    generate_button = tkinter.Button(window, text="생성하기", overrelief="solid", width=15, command=cert_generator)
    generate_button.pack()
    exit_button = tkinter.Button(window, text="종료하기", overrelief="solid", width=15, command=exit)
    exit_button.pack()

    window.mainloop()